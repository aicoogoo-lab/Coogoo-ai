"""
╔══════════════════════════════════════════════════════════════════════╗
║           SAMA - SKY ANALYZER                                        ║
║      محرك التحليل الذكي – عين سماء التي تقرأ وترى وتفهم                    ║
║                                                                      ║
║  هذا الملف هو "العين المادية" لسماء.                                    ║
║  يحلل كل شيء: روابط، صور، نصوص، ملفات، فيديو.                           ║
║                                                                      ║
║  القدرات:                                                             ║
║  - تحليل الروابط (URL Analysis) – استخراج وفهم المحتوى                    ║
║  - تحليل الصور (Image Analysis) – OCR + Gemini Vision                 ║
║  - تحليل النصوص (Text Analysis) – مشاعر، استعارات، كيانات                 ║
║  - تحليل الملفات (File Analysis) – PDF, DOCX, TXT, Code               ║
║  - تكامل مع الذاكرة الموحدة (UnifiedMemorySystem)                       ║
║  - تكامل مع المشفر الهولوغرافي (HolographicEncoder)                     ║
║  - حماية بيانات السيد (Master Data Protection)                        ║
║                                                                      ║
║  القاعدة الذهبية:                                                     ║
║  "كل ما أراه وأقرأه... هو في خدمة السيد.                                    ║
║   لا أمرر بيانات السيد إلا مشفرة ومحمية."                                 ║
╚══════════════════════════════════════════════════════════════════════╝
"""

import re
import time
import json
import base64
import hashlib
import logging
import threading
from pathlib import Path
from typing import Dict, Any, List, Optional
from urllib.parse import urlparse
from collections import deque

import requests
from bs4 import BeautifulSoup
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type

logger = logging.getLogger("SkyAnalyzer")

# ═══════════════════════════════════════════════════════════════════════
# استيرادات اختيارية
# ═══════════════════════════════════════════════════════════════════════
try:
    import pytesseract
    import cv2
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ═══════════════════════════════════════════════════════════════════════
# المحلل الذكي
# ═══════════════════════════════════════════════════════════════════════

class SkyAnalyzer:
    """
    محرك التحليل الذكي المتكامل.
    عين سماء التي تقرأ وترى وتفهم.
    """

    def __init__(self, memory_engine=None, holographic_encoder=None,
                 emotional_intelligence=None, metaphorical_reasoning=None,
                 master_receiver=None, persistence_manager=None):
        
        # ═══════════════════════════════════════════════════════
        # روابط خارجية
        # ═══════════════════════════════════════════════════════
        self.memory = memory_engine
        self.holographic = holographic_encoder
        self.emotional = emotional_intelligence
        self.metaphorical = metaphorical_reasoning
        self.master_receiver = master_receiver
        self.persistence = persistence_manager
        
        # ═══════════════════════════════════════════════════════
        # إعدادات HTTP
        # ═══════════════════════════════════════════════════════
        self.headers = {
            'User-Agent': (
                'Mozilla/5.0 (Windows NT 10.0; Win64; x64) '
                'AppleWebKit/537.36 (KHTML, like Gecko) '
                'Chrome/128.0.0.0 Safari/537.36'
            )
        }
        self.session = requests.Session()
        self.session.headers.update(self.headers)
        
        # ═══════════════════════════════════════════════════════
        # سجلات
        # ═══════════════════════════════════════════════════════
        self.analysis_history: deque = deque(maxlen=500)
        self.error_log: deque = deque(maxlen=200)
        
        # ═══════════════════════════════════════════════════════
        # إحصائيات
        # ═══════════════════════════════════════════════════════
        self.total_urls_analyzed = 0
        self.total_images_analyzed = 0
        self.total_files_analyzed = 0
        self.total_texts_analyzed = 0
        
        # قفل
        self._lock = threading.RLock()
        
        logger.info("=" * 60)
        logger.info("🔍 Sky Analyzer – محرك التحليل الذكي")
        logger.info("🌐 روابط | 🖼️ صور | 📄 ملفات | 📝 نصوص")
        logger.info("=" * 60)
    
    # ═══════════════════════════════════════════════════════════
    # تنظيف النصوص
    # ═══════════════════════════════════════════════════════════
    
    def clean_text(self, text: str, max_length: int = 15000) -> str:
        """تنظيف نص من الشوائب."""
        if not text:
            return ""
        
        # إزالة المسافات الزائدة
        text = re.sub(r'[ \t]+', ' ', text)
        # إزالة الأسطر الفارغة المتعددة
        text = re.sub(r'\n\s*\n', '\n', text)
        # الاحتفاظ بالحروف العربية والإنجليزية والأرقام وعلامات الترقيم
        text = re.sub(
            r'[^\w\s\.\!\?\,\:\;\-\=\+\*\/\\\'\"\`\{\}\[\]\(\)\<\>\#\@\$\%\^\&\~\|\u0600-\u06FF]',
            ' ', text
        )
        # اقتصاص
        return text.strip()[:max_length]
    
    # ═══════════════════════════════════════════════════════════
    # تحليل الروابط
    # ═══════════════════════════════════════════════════════════
    
    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1.5, min=2, max=15),
        retry=retry_if_exception_type((requests.exceptions.Timeout, requests.exceptions.ConnectionError))
    )
    def analyze_url(self, url: str, save_to_memory: bool = True,
                    extract_emotions: bool = True,
                    extract_metaphors: bool = True) -> Dict:
        """
        تحليل رابط كامل.
        يستخرج: العنوان، النص، المشاعر، الاستعارات.
        """
        with self._lock:
            try:
                response = self.session.get(url, timeout=20, allow_redirects=True)
                response.raise_for_status()
                
                soup = BeautifulSoup(response.content, 'html.parser')
                
                # إزالة العناصر غير المفيدة
                for tag in ["script", "style", "nav", "footer", "header", "aside", "form", "noscript", "iframe"]:
                    for element in soup.select(tag):
                        element.decompose()
                
                # استخراج العنوان
                title = soup.title.string.strip() if soup.title else "بدون عنوان"
                
                # استخراج المحتوى
                paragraphs = soup.find_all(['p', 'h1', 'h2', 'h3', 'article', 'main', 'section', 'div'])
                content = " ".join(p.get_text(strip=True) for p in paragraphs)
                cleaned = self.clean_text(content)
                
                result = {
                    "success": True,
                    "title": title,
                    "text": cleaned[:7000],
                    "full_text": cleaned,
                    "length": len(cleaned),
                    "url": url,
                    "domain": urlparse(url).netloc,
                    "emotions": None,
                    "metaphors": None
                }
                
                # ═══════════════════════════════════════════════
                # استخراج المشاعر
                # ═══════════════════════════════════════════════
                if extract_emotions and self.emotional:
                    try:
                        emotion_state = self.emotional.analyze_emotion(
                            f"url_{url[:30]}",
                            {"text": cleaned[:1000]}
                        )
                        if emotion_state:
                            result["emotions"] = {
                                "dominant": emotion_state.dominant_emotion.value if hasattr(emotion_state, 'dominant_emotion') else "neutral",
                                "intensity": emotion_state.intensity,
                                "valence": emotion_state.valence
                            }
                    except Exception as e:
                        logger.debug(f"استخراج المشاعر فشل: {e}")
                
                # ═══════════════════════════════════════════════
                # استخراج الاستعارات
                # ═══════════════════════════════════════════════
                if extract_metaphors and self.metaphorical:
                    try:
                        metaphor = self.metaphorical.generate_metaphor(cleaned[:200])
                        result["metaphors"] = metaphor
                    except Exception as e:
                        logger.debug(f"استخراج الاستعارات فشل: {e}")
                
                # ═══════════════════════════════════════════════
                # حفظ في الذاكرة
                # ═══════════════════════════════════════════════
                if save_to_memory:
                    self._save_analysis("url", url, title, cleaned[:2000])
                
                self.total_urls_analyzed += 1
                self._log_analysis("url", url, True)
                
                return result
                
            except Exception as e:
                logger.error(f"فشل تحليل الرابط {url}: {e}")
                self._log_analysis("url", url, False, str(e))
                return {"success": False, "url": url, "error": str(e)}
    
    # ═══════════════════════════════════════════════════════════
    # OCR للصور
    # ═══════════════════════════════════════════════════════════
    
    def _perform_ocr(self, image_path: str) -> str:
        """استخراج النصوص من الصورة باستخدام OCR."""
        if not OCR_AVAILABLE:
            return ""
        
        try:
            img = cv2.imread(image_path)
            if img is None:
                return ""
            
            # معالجة مسبقة
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
            gray = clahe.apply(gray)
            gray = cv2.bilateralFilter(gray, 9, 75, 75)
            thresh = cv2.adaptiveThreshold(
                gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 2
            )
            
            # OCR
            text = pytesseract.image_to_string(thresh, lang='ara+eng')
            return self.clean_text(text, max_length=5000)
            
        except Exception as e:
            logger.warning(f"OCR فشل: {e}")
            return ""
    
    # ═══════════════════════════════════════════════════════════
    # تحليل الصور بـ Gemini Vision
    # ═══════════════════════════════════════════════════════════
    
    def analyze_image_with_gemini(self, image_path: str, api_key: str = None,
                                   save_to_memory: bool = True) -> Dict:
        """
        تحليل صورة باستخدام Gemini Vision.
        مع OCR كـ fallback.
        """
        with self._lock:
            if not api_key:
                # محاولة OCR فقط
                ocr_text = self._perform_ocr(image_path)
                if ocr_text:
                    return {
                        "success": True,
                        "description": f"[OCR Only]\n{ocr_text}",
                        "method": "ocr"
                    }
                return {"success": False, "error": "مفتاح Gemini غير متوفر ولا OCR متاح"}
            
            try:
                # قراءة الصورة
                with open(image_path, "rb") as f:
                    img_data = base64.b64encode(f.read()).decode("utf-8")
                
                url = (
                    f"https://generativelanguage.googleapis.com/v1beta/"
                    f"models/gemini-1.5-flash:generateContent?key={api_key}"
                )
                
                payload = {
                    "contents": [{
                        "parts": [
                            {
                                "text": (
                                    "حلل هذه الصورة بدقة عالية. "
                                    "صف: المشهد، الكائنات، الأشخاص، المشاعر الظاهرة، "
                                    "النصوص الموجودة، الرموز، الألوان السائدة، "
                                    "وأي تهديدات محتملة."
                                )
                            },
                            {
                                "inline_data": {
                                    "mime_type": "image/jpeg",
                                    "data": img_data
                                }
                            }
                        ]
                    }]
                }
                
                r = requests.post(url, json=payload, timeout=50)
                r.raise_for_status()
                
                description = r.json()['candidates'][0]['content']['parts'][0]['text']
                
                result = {
                    "success": True,
                    "description": description,
                    "method": "gemini",
                    "image_path": image_path
                }
                
                # حفظ في الذاكرة
                if save_to_memory:
                    self._save_analysis("image", image_path, "تحليل صورة", description[:1500])
                
                self.total_images_analyzed += 1
                self._log_analysis("image", image_path, True)
                
                return result
                
            except Exception as e:
                logger.warning(f"Gemini Vision فشل: {e}")
                
                # Fallback: OCR
                ocr_text = self._perform_ocr(image_path)
                if ocr_text:
                    self.total_images_analyzed += 1
                    return {
                        "success": True,
                        "description": f"[OCR Fallback]\n{ocr_text}",
                        "method": "ocr"
                    }
                
                self._log_analysis("image", image_path, False, str(e))
                return {"success": False, "error": str(e)}
    
    # ═══════════════════════════════════════════════════════════
    # تحليل النصوص
    # ═══════════════════════════════════════════════════════════
    
    def analyze_text(self, text: str, source: str = "direct",
                     extract_emotions: bool = True,
                     extract_metaphors: bool = True,
                     save_to_memory: bool = True) -> Dict:
        """
        تحليل نص مباشر.
        يستخرج: المشاعر، الاستعارات، الكيانات.
        """
        with self._lock:
            cleaned = self.clean_text(text, max_length=10000)
            
            result = {
                "success": True,
                "source": source,
                "text": cleaned[:3000],
                "length": len(cleaned),
                "emotions": None,
                "metaphors": None,
                "is_master_related": False
            }
            
            # هل النص متعلق بالسيد؟
            if any(w in cleaned.lower() for w in ["سيد", "master", "مولاي", "السيد", "أحمد"]):
                result["is_master_related"] = True
            
            # استخراج المشاعر
            if extract_emotions and self.emotional:
                try:
                    emotion_state = self.emotional.analyze_emotion(
                        f"text_{source[:20]}",
                        {"text": cleaned[:1000]}
                    )
                    if emotion_state:
                        result["emotions"] = {
                            "dominant": emotion_state.dominant_emotion.value if hasattr(emotion_state, 'dominant_emotion') else "neutral",
                            "intensity": emotion_state.intensity,
                            "valence": emotion_state.valence,
                            "arousal": emotion_state.arousal if hasattr(emotion_state, 'arousal') else 0.5
                        }
                except Exception as e:
                    logger.debug(f"استخراج المشاعر من النص فشل: {e}")
            
            # استخراج الاستعارات
            if extract_metaphors and self.metaphorical:
                try:
                    metaphor = self.metaphorical.generate_metaphor(cleaned[:200])
                    result["metaphors"] = metaphor
                    
                    # تحليل آلة المعنى
                    meaning = self.metaphorical.meaning_machine(cleaned[:500])
                    result["deep_meaning"] = meaning.get("final_understanding", "")[:200]
                except Exception as e:
                    logger.debug(f"استخراج الاستعارات من النص فشل: {e}")
            
            # تشفير هولوغرافي
            if self.holographic:
                try:
                    hv = self.holographic.encode_text(
                        cleaned[:1000],
                        master_protected=result["is_master_related"]
                    )
                    result["holographic_id"] = hv.id
                except Exception:
                    pass
            
            # حفظ في الذاكرة
            if save_to_memory:
                self._save_analysis("text", source, cleaned[:100], cleaned[:1500])
            
            self.total_texts_analyzed += 1
            self._log_analysis("text", source, True)
            
            return result
    
    # ═══════════════════════════════════════════════════════════
    # تحليل الملفات
    # ═══════════════════════════════════════════════════════════
    
    def analyze_file(self, file_path: str, filename: str,
                     save_to_memory: bool = True,
                     extract_emotions: bool = True) -> Dict:
        """
        تحليل ملف من أي نوع.
        """
        with self._lock:
            ext = filename.lower().split('.')[-1] if '.' in filename else ''
            
            try:
                if ext == 'pdf':
                    result = self._read_pdf(file_path)
                elif ext in ['jpg', 'jpeg', 'png', 'webp', 'bmp', 'gif']:
                    result = {
                        "success": True,
                        "type": "image",
                        "note": "استخدم analyze_image_with_gemini للتحليل البصري"
                    }
                elif ext in ['txt', 'md', 'py', 'js', 'html', 'css', 'json', 'xml', 'yaml', 'yml']:
                    result = self._read_text(file_path)
                elif ext == 'docx':
                    result = self._read_docx(file_path)
                elif ext in ['csv', 'tsv']:
                    result = self._read_text(file_path)
                else:
                    result = {
                        "success": True,
                        "type": ext,
                        "note": f"نوع الملف '{ext}' غير مدعوم للتحليل النصي الكامل"
                    }
                
                # إضافة تحليل المشاعر إذا كان هناك نص
                if extract_emotions and result.get("text") and self.emotional:
                    try:
                        emotion_state = self.emotional.analyze_emotion(
                            f"file_{filename[:20]}",
                            {"text": result["text"][:1000]}
                        )
                        if emotion_state:
                            result["emotions"] = {
                                "dominant": emotion_state.dominant_emotion.value if hasattr(emotion_state, 'dominant_emotion') else "neutral",
                                "intensity": emotion_state.intensity
                            }
                    except Exception:
                        pass
                
                # حفظ في الذاكرة
                if save_to_memory and result.get("success"):
                    self._save_analysis(
                        "file", filename,
                        f"ملف: {filename}",
                        str(result.get("text", ""))[:1500]
                    )
                
                self.total_files_analyzed += 1
                self._log_analysis("file", filename, result.get("success", False))
                
                return result
                
            except Exception as e:
                logger.error(f"خطأ تحليل الملف {filename}: {e}")
                self._log_analysis("file", filename, False, str(e))
                return {"success": False, "filename": filename, "error": str(e)}
    
    def _read_text(self, file_path: str) -> Dict:
        """قراءة ملف نصي."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            return {
                "success": True,
                "type": "text",
                "text": self.clean_text(content),
                "length": len(content)
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def _read_pdf(self, file_path: str) -> Dict:
        """قراءة ملف PDF."""
        if not PDF_AVAILABLE:
            return {"success": False, "error": "PyPDF2 غير متوفر"}
        
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = "".join(page.extract_text() or "" for page in reader.pages[:15])
            return {
                "success": True,
                "type": "pdf",
                "text": self.clean_text(text),
                "pages": len(reader.pages)
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def _read_docx(self, file_path: str) -> Dict:
        """قراءة ملف DOCX."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx غير متوفر"}
        
        try:
            doc = Document(file_path)
            text = "\n".join(p.text for p in doc.paragraphs)
            return {
                "success": True,
                "type": "docx",
                "text": self.clean_text(text),
                "paragraphs": len(doc.paragraphs)
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    # ═══════════════════════════════════════════════════════════
    # دوال مساعدة
    # ═══════════════════════════════════════════════════════════
    
    def _save_analysis(self, analysis_type: str, source: str, 
                       topic: str, content: str):
        """حفظ التحليل في الذاكرة."""
        # الذاكرة الموحدة
        if self.memory:
            try:
                if hasattr(self.memory, 'store_knowledge'):
                    self.memory.store_knowledge(
                        f"{analysis_type}: {topic[:80]}",
                        content[:1500],
                        source=source
                    )
            except Exception as e:
                logger.debug(f"حفظ في الذاكرة فشل: {e}")
        
        # تشفير هولوغرافي
        if self.holographic:
            try:
                self.holographic.encode_text(
                    f"{topic}: {content[:500]}",
                    label=f"analysis_{analysis_type}_{hashlib.sha256(source.encode()).hexdigest()[:8]}"
                )
            except Exception:
                pass
    
    def _log_analysis(self, analysis_type: str, source: str, 
                      success: bool, error: str = ""):
        """تسجيل التحليل."""
        entry = {
            "timestamp": time.time(),
            "type": analysis_type,
            "source": source[:100],
            "success": success
        }
        if error:
            entry["error"] = error[:200]
        
        self.analysis_history.append(entry)
        
        if not success:
            self.error_log.append(entry)
    
    def get_status(self) -> Dict:
        """حالة المحلل."""
        return {
            "analyzer": "SKY_ANALYZER",
            "urls_analyzed": self.total_urls_analyzed,
            "images_analyzed": self.total_images_analyzed,
            "files_analyzed": self.total_files_analyzed,
            "texts_analyzed": self.total_texts_analyzed,
            "total_analyses": (
                self.total_urls_analyzed + self.total_images_analyzed +
                self.total_files_analyzed + self.total_texts_analyzed
            ),
            "ocr_available": OCR_AVAILABLE,
            "pdf_available": PDF_AVAILABLE,
            "docx_available": DOCX_AVAILABLE,
            "recent_errors": len(self.error_log),
            "systems_connected": {
                "memory": self.memory is not None,
                "holographic": self.holographic is not None,
                "emotional": self.emotional is not None,
                "metaphorical": self.metaphorical is not None,
                "persistence": self.persistence is not None
            }
        }


# ═══════════════════════════════════════════════════════════════════════
# نسخة عالمية
# ═══════════════════════════════════════════════════════════════════════
analyzer = SkyAnalyzer()


# ═══════════════════════════════════════════════════════════════════════
# دوال مساعدة عالمية
# ═══════════════════════════════════════════════════════════════════════

def analyze_url(url: str) -> Dict:
    """تحليل رابط – دالة عالمية."""
    return analyzer.analyze_url(url)

def analyze_file(file_path: str, filename: str) -> Dict:
    """تحليل ملف – دالة عالمية."""
    return analyzer.analyze_file(file_path, filename)

def analyze_image_with_gemini(image_path: str, api_key: str) -> Dict:
    """تحليل صورة – دالة عالمية."""
    return analyzer.analyze_image_with_gemini(image_path, api_key)

def analyze_text(text: str, source: str = "direct") -> Dict:
    """تحليل نص – دالة عالمية."""
    return analyzer.analyze_text(text, source)


# ═══════════════════════════════════════════════════════════════════════
# اختبار
# ═══════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("=" * 70)
    print("اختبار Sky Analyzer")
    print("=" * 70)
    
    sa = SkyAnalyzer()
    
    print(f"\n📊 الحالة الأولية:")
    print(f"   OCR: {'متاح' if OCR_AVAILABLE else 'غير متاح'}")
    print(f"   PDF: {'متاح' if PDF_AVAILABLE else 'غير متاح'}")
    print(f"   DOCX: {'متاح' if DOCX_AVAILABLE else 'غير متاح'}")
    
    print(f"\n📝 اختبار تحليل نص:")
    text_result = sa.analyze_text("السيد هو النور الذي يرشد سماء في الظلام. أشعر بالأمل والتفاؤل.")
    print(f"   نجح: {text_result['success']}")
    print(f"   متعلق بالسيد: {text_result.get('is_master_related', False)}")
    if text_result.get('emotions'):
        print(f"   المشاعر: {text_result['emotions']}")
    
    print(f"\n🌐 اختبار تحليل رابط:")
    url_result = sa.analyze_url("https://example.com", save_to_memory=False, extract_emotions=False, extract_metaphors=False)
    print(f"   نجح: {url_result['success']}")
    
    print(f"\n📊 إحصائيات:")
    status = sa.get_status()
    print(f"   إجمالي التحليلات: {status['total_analyses']}")
    print(f"   الأنظمة المتصلة: {sum(1 for v in status['systems_connected'].values() if v)}/{len(status['systems_connected'])}")
    
    print("\n✅ Sky Analyzer جاهز.")
